from ast import main
from docx import Document
from indexed import IndexedOrderedDict
from docx.shared import Pt
from docx.shared import RGBColor
from utils.valve2 import Valve2
from utils.valve3 import Valve3
from utils.line import Line
import os

class DocWriter():
    def __init__(self, name="MyDoc"):
        ### note python "self" is like "this" in c#
        ### "__init__" is a constructor so that different options can be set.
        self.name = name
        self.doc = Document()
        run = self.doc.add_heading("", 1).add_run()
        font = run.font
        font.name = "Times New Roman"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x42,0x42,0x42)
        run.add_text(self.name)

        
    def makeSection(self, name, instruction = None):
        paragraph = self.doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.space_before = Pt(1)
        section_name = paragraph.add_run(name)
        section_name.font.bold = True
        section_name.font.name = "Times New Roman"
        section_name.font.size = Pt(11)
        prompt = paragraph.add_run(instruction)
        prompt.font.name = "Times New Roman"
        prompt.font.size = Pt(11)
        return paragraph


    def save(self, filename= "PrintedRoute.docx"):
        self.doc.save(filename)

    def buildDocument(self, simple_route, dvi_route, pits):
        used_pits = IndexedOrderedDict()
        directly_used_pits = IndexedOrderedDict()
        used_jumpers = IndexedOrderedDict()
        used_lines = []
        #directly used lines, ie. not dvi
        directly_used_lines = []
        for component in simple_route:
            if component.pit and component.pit in pits:
                
                directly_used_pits[component.pit] = pits[component.pit] 
            if (type(component) == Line):
                directly_used_lines.append(component)            
        for component in dvi_route:
            if component.pit and component.pit in pits:
                used_pits[component.pit] = pits[component.pit] 
                used_pits[component.pit].add_used_component(component)
            if component.onJumper:
                jumper = (component.pit, component.jumper)
                used_jumpers[jumper] = None
            if (type(component) == Line):
                used_lines.append(component)
        wlps_text = self.makeSection("Route Description: ", "use description for Waste Leak Path Screen")
        wlps_text.add_run("\n")
        
        ###
        sending_tank = used_pits.values()[0].tsr_structure[:-3] + "1" + directly_used_pits.values()[0].tsr_structure[-3:-1]
        receiving_tank = used_pits.values()[-1].tsr_structure[:-3] + "1" + directly_used_pits.values()[-1].tsr_structure[-3:-1]
        wlps_text.add_run(f"Waste from tank {sending_tank} will be transferred using {simple_route[0]}, routed through ")
        for pit, line in zip(directly_used_pits, directly_used_lines):
            wlps_text.add_run(f"{pit} jumpers, ")
            wlps_text.add_run(f"{line.ein[-6:]}, ")    
        wlps_text.add_run(f" finally discharging into tank {receiving_tank}'s head space through the drop leg at {simple_route[-1]}.")
        procedure_development_data = self.makeSection("Procedure Development Data")
        heaterEINs = self.makeSection("Section 5.5.3 heaters: " ,"Replace existing data with the following:")
        for pit in used_pits.values():
            for heater in pit.in_pit_heaters:
                heaterEINs.add_run("\n")
                heaterEINs.add_run(heater)
                heaterEINs.add_run("\t \t")      
                heaterEINs.add_run(pit.pit_nace)
        pits5179 = self.makeSection("Steps 5.17.9: ","Replace existing data with the following:")
        for pit in used_pits.values():
            pits5179.add_run("\n")
            pits5179.add_run(pit.drain_seal_location)
        checklist1 = self.makeSection("Checklist 1: ","Replace list with:")
        for jumper in used_jumpers:
            checklist1.add_run("\n")
            checklist1.add_run(jumper[0])
            checklist1.add_run("\t \t \t")
            checklist1.add_run(f"Jumper: {jumper[1]} ")
        checklist3 = self.makeSection("Checklist 3: Transfer Valving","")
        for pit in used_pits.values():
            checklist3.add_run("\n")
            checklist3.add_run(pit.pit_nace).bold = True
            checklist3.add_run(" Tank Farm").bold = True
            for component in pit.components:
                if (type(component) == Valve3 or type(component) == Valve2 ):
                    checklist3.add_run("\n")
                    checklist3.add_run(component.EIN())
                    checklist3.add_run("\t \t")
                    checklist3.add_run(component.position)
                    if (component.dvi_used == "YES" or component.dvi_used == "POS"):
                        checklist3.add_run("\t ")
                        checklist3.add_run("(Mark as DVI)").bold = True
            checklist3.add_run("\n")
            checklist3.add_run(f"Confirm open route: ({pit.pit_nace}) ").bold = True
            checklist3.add_run("\n")
            checklist3.add_run(f"FROM \n").bold = True
            checklist3.add_run(pit.components[0].field_label)
            checklist3.add_run(f"\nTO\n").bold = True
            if pit.components[-1].field_label:
                checklist3.add_run(pit.components[-1].field_label)
            else:
                checklist3.add_run(pit.components[-1].ein)
            checklist3.add_run("\n")
        checklist4 = self.makeSection("Checklist 4: Checklist 4 - Flush Transfer Route to Transfer Pump Valving","")
        checklist5 = self.makeSection("Checklist 5: Checklist 5 - Flush Transfer Route to Receiving Tank Valving","")
        checklist6 = self.makeSection("Checklist 6: Return to Transfer Valving","")
        checklist7LD = self.makeSection("Checklist 7 - Tank pit/Structure Leak Detection")
        checklist7TF = self.makeSection("Checklist 7 - TFSPS Temperature Equipment Checks")
        for pit in used_pits.values():
            for tfsps , pmid in zip(pit.tfsps_transmitters, pit.tfsps_pmids):
                checklist7TF.add_run("\n")
                checklist7TF.add_run(tfsps)
                checklist7TF.add_run("\t \t")
                checklist7TF.add_run(pmid)
        checklist7D = self.makeSection("Checklist 7 - Drain Seal Assemblies:")
        for pit in used_pits.values():
            checklist7D.add_run("\n")
            checklist7D.add_run(pit.drain_seal_location)
            checklist7D.add_run("\t \t")
            checklist7D.add_run(pit.drain_seal_position)
        checklist7N = self.makeSection("Checklist 7 - NACE Inspection:")
        for pit in used_pits.values():
            checklist7N.add_run("\n")
            checklist7N.add_run(pit.pit_nace)
            checklist7N.add_run("\t \t")
            checklist7N.add_run(pit.pit_nace_pmid)
        # route_list = self.makeSection("SECD Route List: ")
        # for node in dvi_route:
        #     if node.show:
        #         route_list.add_run("\n")
        #         route_list.add_run(node.EIN())



